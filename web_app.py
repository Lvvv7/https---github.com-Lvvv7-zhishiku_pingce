from flask import Flask, render_template_string, request, jsonify, send_file, render_template
import os
import sqlite3
import time
import docx
import requests
import json
from openpyxl import Workbook
from werkzeug.utils import secure_filename
import markdown as md
from markupsafe import Markup

app = Flask(__name__)

# 数据库初始化
DB_PATH = os.path.join(os.path.dirname(__file__), 'query_logs.db')
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        url TEXT,
        question TEXT,
        resp TEXT,
        auth TEXT,
        timestamp TEXT
    )''')
    conn.commit()
    conn.close()
init_db()

def save_log(url, question, resp, auth):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('INSERT INTO logs (url, question, resp, auth, timestamp) VALUES (?, ?, ?, ?, ?)',
              (url, question, resp, auth, time.strftime('%Y-%m-%d %H:%M:%S')))
    conn.commit()
    conn.close()
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 限制文件大小为16MB

# 确保上传文件夹存在
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

app.jinja_env.filters['markdown'] = lambda text: Markup(md.markdown(text or "", extensions=['extra', 'tables']))

## 前端页面已迁移至 templates/index.html 和 static/

def process_document(doc_path):
    try:
        doc = docx.Document(doc_path)
        all_content = []

        # 处理段落
        for para in doc.paragraphs:
            all_content.append(para.text)

        # 处理表格
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_content = [cell.text.strip() for cell in row.cells]
                table_content.append(row_content)
            
            # 转换为Markdown表格
            if table_content:
                markdown_table = "| " + " | ".join(table_content[0]) + " |\n"
                markdown_table += "| " + " | ".join(["---"] * len(table_content[0])) + " |\n"
                for row in table_content[1:]:
                    markdown_table += "| " + " | ".join(row) + " |\n"
                all_content.append(markdown_table)

        full_content = "\n\n".join(all_content)

        # 调用API处理内容
        url = "https://api.siliconflow.cn/v1/chat/completions"
        payload = {
            "model": "Qwen/Qwen3-32B",
            "stream": False,
            "max_tokens": 2000,
            "enable_thinking": True,
            "thinking_budget": 4096,
            "min_p": 0.05,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "frequency_penalty": 0.5,
            "n": 1,
            "stop": [],
            "messages": [
                {
                    "role": "user",
                    "content": full_content
                },
                {
                    "role": "system",
                    "content": "你是一个文档信息整理智能助手，帮助用户将问文档的信息整理成20条问题和该问题的答案。答案前的字符为'答案：'，只需输出问题和答案，其他无关的信息不要输出。"
                }
            ]
        }
        headers = {
            "Authorization": "Bearer sk-rmvrrwmtrqqankpkwgnpyuvjvsrlbbgbtpwxdsehjycikkfd",
            "Content-Type": "application/json"
        }

        response = requests.post(url, json=payload, headers=headers)
        output_data = response.json()
        output_data = output_data["choices"][0]["message"]["content"]

        # 创建Excel文件
        wb = Workbook()
        ws = wb.active
        ws['A1'] = '问题'
        ws['B1'] = '答案'

        lines = output_data.strip().split("\n\n")
        for index, line in enumerate(lines, start=2):
            parts = line.split("答案：")
            if len(parts) == 2:
                question = parts[0].strip()
                answer = parts[1].strip()
                ws.cell(row=index, column=1, value=question)
                ws.cell(row=index, column=2, value=answer)

        output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'output.xlsx')
        wb.save(output_path)
        return True
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return False

@app.route('/')
def index():
    from flask import render_template
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': '没有上传文件'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': '未选择文件'})
    
    if not file.filename.endswith('.docx'):
        return jsonify({'success': False, 'error': '请上传.docx格式的文件'})
    
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        if process_document(filepath):
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': '处理文件时发生错误'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/download')
def download():
    try:
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'output.xlsx')
        return send_file(output_path, as_attachment=True, download_name='处理结果.xlsx')
    except Exception as e:
        return jsonify({'success': False, 'error': '下载文件时发生错误'})

@app.route('/result')
def result():
    import openpyxl
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'output.xlsx')
    table = []
    if os.path.exists(output_path):
        wb = openpyxl.load_workbook(output_path)
        ws = wb.active
        for row in ws.iter_rows(min_row=2, values_only=True):
            table.append(row)
    return render_template('result.html', table=table)

@app.route('/workflow', methods=['GET', 'POST'])
def workflow():
    import openpyxl
    questions = []
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'output.xlsx')
    if os.path.exists(output_path):
        wb = openpyxl.load_workbook(output_path)
        ws = wb.active
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row and row[0]:
                questions.append(str(row[0]))
    results = []
    if request.method == 'POST':
        api_urls = request.form.get('api_urls', '').strip().splitlines()
        auth = request.form.get('auth', '').strip()
        for url in api_urls:
            url = url.strip()
            if not url:
                continue
            for q in questions:
                try:
                    payload = {
                        "inputs": {},
                        "response_mode": "blocking",
                        "auto_generate_name": True,
                        "query": q,
                        "user": "a",
                        "conversation_id": "884ea56b-9663-49aa-9538-517aca3fa03e"
                    }
                    headers = {
                        "Authorization": auth,
                        "Content-Type": "application/json"
                    }
                    resp = requests.post(url, json=payload, headers=headers, timeout=60)
                    try:
                        resp_json = resp.json()
                        answer = resp_json.get("answer", "无answer字段")
                        resp_text = json.dumps(answer, ensure_ascii=False, indent=2)
                    except Exception:
                        resp_text = resp.text
                except Exception as e:
                    resp_text = f'请求失败: {e}'
                results.append({'url': url, 'question': q, 'resp': resp_text})
                save_log(url, q, resp_text, auth)
    return render_template('workflow.html', results=results, questions=questions)

@app.route('/workflow_api', methods=['POST'])
def workflow_api():
    url = request.form.get('url', '').strip()
    auth = request.form.get('auth', '').strip()
    question = request.form.get('question', '').strip()
    if not url or not auth or not question:
        return jsonify({'error': '参数缺失'}), 400
    try:
        payload = {
            "inputs": {},
            "response_mode": "blocking",
            "auto_generate_name": True,
            "query": question,
            "user": "a",
            "conversation_id": "de347d53-6a39-4e68-b83b-2c01a09c0ca8"
        }
        headers = {
            "Authorization": auth,
            "Content-Type": "application/json"
        }
        resp = requests.post(url, json=payload, headers=headers, timeout=60)
        try:
            resp_json = resp.json()
            answer = resp_json.get("answer", "无answer字段")
            resp_text = answer
        except Exception:
            resp_text = resp.text
    except Exception as e:
        resp_text = f'请求失败: {e}'
    save_log(url, question, resp_text, auth)
    return jsonify({'url': url, 'question': question, 'resp': resp_text})

# 查询历史记录页面
@app.route('/logs')
def logs():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('SELECT url, question, resp, auth, timestamp FROM logs ORDER BY id DESC')
    logs = c.fetchall()
    conn.close()
    return render_template('logs.html', logs=logs)

if __name__ == '__main__':
    app.run(debug=True, port=5000)
